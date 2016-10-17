# -*- coding: utf-8 -*-
from bs4 import BeautifulSoup, element
import smtplib
import json
from docx import Document
from docx.shared import RGBColor, Pt
import os
import codecs
import re
import mimetypes
from email import encoders
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

from logger import logger
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC

from config import *

class ArticleCrawler(object):
    def __init__(self, url=WORKING_URL, output_dir=OUTPUT_DIRECTORY, output_fn=OUTPUT_FILENAME,
                 datafile=USER_DATA_FILE, smtp_server=SMTP_SERVER, smtp_port=SMTP_PORT):
        self.url = url
        self.output_dir = output_dir
        self.output_fn = output_fn
        self.datafile = datafile
        self.smtp_server = smtp_server
        self.smtp_port = smtp_port

        with open(USER_DATA_FILE) as f:
            self.old_articles = json.load(f)

    def get_web_content(self, url=None, present_tag="class", present_element="main_col"):
        target_url = url if url else self.url
        tag = By.ID if present_tag == 'id' else By.CLASS_NAME
        driver = webdriver.Chrome(CHROME_DRIVER)
        driver.get(target_url)
        try:
            element = WebDriverWait(driver, 10).until(
                EC.presence_of_element_located((tag, present_element))
            )
        finally:
            soup = BeautifulSoup(driver.page_source, "html.parser")
            driver.quit()
            return soup

    def get_articles(self, url=None):
        content = self.get_web_content(url=url)
        return [(item.string.strip(), WORKING_DOMAIN + item['href']) for item in content.find_all(class_="question_link")]

    def get_new_articles(self, url=None):
        articles = self.get_articles(url=url)
        return [item for item in articles if item[0] not in self.old_articles and TARGET_TITLE in item[0]]

    def get_article_content(self, url):
        content = self.get_web_content(url, "id", "js_content")
        return content

    def find_initial_p(self, all_p, start_num):
        init = 0
        for i, p in enumerate(all_p):
            if self.is_title(p) and p.strong.contents[0].startswith(start_num):
                init = i
                break
        return init

    def output_to_html(self, title, content):
        start_num, end_num = title.replace(TARGET_TITLE, "").replace(" ", "").split("-")
        all_p = content.find_all("p")
        init = self.find_initial_p(all_p, start_num)

        output_filename = title + u".html"
        with codecs.open(os.path.join(self.output_dir, output_filename), 'w', encoding="utf-8") as f:
            beginning_html = '<!DOCTYPE html><html>' \
                             '<meta http-equiv="Content-Type" content="text/html; charset=utf-8" /><body>'
            ending_html = '</body></html>'

            f.write(beginning_html)

            last_one = False
            for p_tag in all_p[init:]:
                if self.is_title(p_tag) and p_tag.strong.contents[0].startswith(end_num):
                    last_one = True

                f.write(unicode(p_tag))

                if last_one and p_tag.find("br"):
                    break

            f.write(ending_html)

        self.update_downloaded_list(title)
        return output_filename

    def output_to_docx(self, title, content):
        start_num, end_num = title.replace(TARGET_TITLE, "").replace(" ", "").split("-")
        all_p = content.find_all("p")
        init = self.find_initial_p(all_p, start_num)
        docx = Document()

        last_one = False
        for p_tag in all_p[init:]:
            is_title = self.is_title(p_tag)
            if is_title and p_tag.strong.contents[0].startswith(end_num):
                last_one = True

            paragraph = docx.add_paragraph()
            self.parse_html_to_docx(p_tag, paragraph, title=is_title)

            if last_one and p_tag.find("br"):
                break

        output_filename = title + u".docx"
        docx.save(os.path.join(self.output_dir, output_filename))
        self.update_downloaded_list(title)
        return output_filename

    def parse_html_to_docx(self, e, paragrash, title=False, red=False):
        if isinstance(e, element.NavigableString):
            if title:
                run = paragrash.add_run(unicode(e))
                run.bold = True
                run.font.size = Pt(18)
            elif red:
                paragrash.add_run(unicode(e)).font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            else:
                paragrash.add_run(unicode(e))
        elif isinstance(e, element.Tag):
            if e.name == "span" and self.is_red(e["style"]):
                for child in e.contents:
                    self.parse_html_to_docx(child, paragrash, title, True)
            elif e.name == "br":
                return
            else:
                for child in e.contents:
                    self.parse_html_to_docx(child, paragrash, title, red)
        else:
            raise Exception("Unknown html element")

    @staticmethod
    def is_red(style):
        if not style:
            return False
        s =re.search("(?<!\-)color:rgb\((\d+),(\d+),(\d+)\)", style.replace(" ", "")) or False
        return s and float(s.group(1)) > 240

    def update_downloaded_list(self, title):
        self.old_articles.append(title)
        with codecs.open(USER_DATA_FILE, 'wb', encoding="utf-8") as f:
            json.dump(self.old_articles, f, indent=2, ensure_ascii=False)

    def run(self, url=None, filetype="docx", email=False):
        articles = self.get_new_articles(url=url)
        attachments = []
        for article in articles:
            if filetype == "html":
                f = self.output_to_html(article[0], self.get_article_content(article[1]))
            else:
                f = self.output_to_docx(article[0], self.get_article_content(article[1]))
            attachments.append(f)
        if email:
            self.email_docx(attachments)

    def get_all(self, ignore_old=False):
        if ignore_old:
            self.old_articles = []
            with codecs.open(USER_DATA_FILE, 'wb', encoding="utf-8") as f:
                json.dump(self.old_articles, f, indent=2, ensure_ascii=False)

        # TODO: retrieve total page number automatically
        total = 1
        for start in range(total):
            url = self.url + "?start={}".format(start*12)
            self.run(url=url, email=True)

    @staticmethod
    def is_title(p):
        return len(p.find_all("strong")) >= 1 and len(p.strong.contents) == 1 and \
               isinstance(p.strong.contents[0], element.NavigableString)

    def is_content(self, p):
        """ Deprecated """
        check = True
        for c in p.contents:
            if not (isinstance(c, unicode) or self.is_emphases(c) or (isinstance(c, element.Tag) and c.name == "em")):
                return False
        return check

    @staticmethod
    def is_emphases(e):
        """ Deprecated """
        return e.name == "span" and u"color: rgb(255, 0, 0)" in e['style']

    def email_docx(self, attachments, email_from=EMAIL_FROM, email_to=EMAIL_TO, subject=EMAIL_SUBJECT,
                   content=EMAIL_CONTENT):
        outer = MIMEMultipart()
        outer['Subject'] = subject
        outer['To'] = ",".join(email_to)
        outer['From'] = email_from

        outer.attach(MIMEText(content, 'plain', 'utf-8'))

        for attachment in attachments:
            file_path = os.path.join(self.output_dir, attachment)
            ctype, encoding = mimetypes.guess_type(file_path)
            if ctype is None or encoding is not None:
                ctype = 'application/octet-stream'
            maintype, subtype = ctype.split('/', 1)

            fp = open(file_path, 'rb')
            msg = MIMEBase(maintype, subtype)
            msg.set_payload(fp.read())
            fp.close()

            # Encode the payload using Base64
            encoders.encode_base64(msg)
            msg.add_header('Content-Disposition', 'attachment', filename=attachment.encode('utf-8'))
            outer.attach(msg)

        print outer
        composed = outer.as_string()

        smtp = smtplib.SMTP_SSL("smtp.gmail.com", 465)
        # smtp.helo()
        # smtp.ehlo()
        # smtp.starttls()
        smtp.login(EMAIL_USERNAME, EMAIL_PASSWORD)
        smtp.sendmail(email_from, email_to, composed)
        smtp.quit()

if __name__ == '__main__':
    crawler = ArticleCrawler()
    crawler.get_all()
    crawler.email_docx(u"【话题语料天天练】151-160.docx")
    # p = BeautifulSoup('<p style="max-width: 100%; min-height: 1em; white-space: pre-wrap; color: rgb(62, 62, 62); line-height: 25px; -webkit-text-stroke-width: initial; text-align: justify; font-family: Avenir; -webkit-text-stroke-color: rgb(0, 0, 0); box-sizing: border-box !important; word-wrap: break-word !important; background-color: rgb(255, 255, 255);"><strong style="max-width: 100%; box-sizing: border-box !important; word-wrap: break-word !important;">141、inside and out 从里到外，彻底</strong></p>', "html.parser")
    # print type(p.find("p")['style'])
    # style = "color: rgb(262, 62, 62); -webkit-text-stroke-color: rgb(62, 62, 62);"
    # import re
    # s = re.search("(?<!\-)color:rgb\((\d+),([0-9]+),([0-9]+)\)", style.replace(" ", "")) or False
    # print s and float(s.group(1)) > 240
    # print mimetypes.guess_type(ur"F:\Jayden\projects\WechatArticleCrawler\【话题语料天天练】151-160.docx")